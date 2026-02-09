from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This places a hyperlink within a paragraph object.
    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

document = Document()

# Set default font to Times New Roman
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Function to apply Times New Roman to runs directly (for safety)
def set_font(run, bold=False, italic=False, size=14):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # This is needed for some Word versions to recognize the font change
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# --- Cover Page ---
title_paragraph = document.add_paragraph()
title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_paragraph.add_run("\n\n\n\nASSIGNMENT\n")
set_font(run, bold=True, size=20)

run = title_paragraph.add_run("ON\n")
set_font(run, size=16)

run = title_paragraph.add_run("THE BHARATIYA NYAYA SANHITA 2023\n\n")
set_font(run, bold=True, size=18)

run = title_paragraph.add_run("Topic:\n")
set_font(run, size=16)

run = title_paragraph.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita\n\n\n")
set_font(run, bold=True, size=18)

# Student Details
details = document.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run("Submitted by:\n")
set_font(run, size=14)
run = details.add_run("Name: [Your Name Here]\n") # Placeholder
set_font(run, bold=True, size=14)
run = details.add_run("Roll No: [Roll No]\n")
set_font(run, bold=True, size=14)
run = details.add_run("Div: [Division]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Submitted to:\n")
set_font(run, size=14)
run = details.add_run("[Professor Name]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Class:\n")
set_font(run, size=14)
run = details.add_run("Third Year BA LLB (Semester VI)\n(A.Y-2025-26)\n\n")
set_font(run, bold=True, size=14)

document.add_page_break()

# --- Main Content ---

# Title
heading = document.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita 2023")
set_font(run, bold=True, size=16)

# Body Layout Settings
def add_body_paragraph(text):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=14)
    return p

# Introduction
add_body_paragraph("Introduction")
intro_text = (
    "The transition from the Indian Penal Code (IPC), 1860 to the Bharatiya Nyaya Sanhita (BNS), 2023 represents more than just a legislative "
    "update; it signifies a fundamental shift in how the Indian state perceives collective violence. For over a century, the crime of mob lynching "
    "existed in a legal grey area, often prosecuted under generic murder statutes that failed to capture the terror of vigilantism. The BNS, effective "
    "from July 2024, finally pierces this veil by institutionalizing specific penalties for mob-led executions. This change is not merely procedural but "
    "is a direct response to the growing demand for accountability in cases where the 'mob' becomes a weapon of prejudice."
)
add_body_paragraph(intro_text)

# Legal Recognition section
add_body_paragraph("Legal Framework and Recognition")
recognition_text = (
    "At the heart of this legal reform is Section 103(2) of the BNS. This provision is innovative because it introduces the concept of five or more "
    "perpetrators acting in concert as a distinct category of offenders when the crime is driven by identity-based hatred. Whether motivated by caste, "
    "language, personal beliefs, or community differences, the law now expressly targets the 'common intention' of a prejudiced group. By doing so, "
    "it dismantles the defense of anonymity that mob members often misused. The statute ensures that when a life is taken by a crowd, the collective "
    "guilt is shared equally among all participants, leaving no room for individuals to hide behind the actions of the many."
)
add_body_paragraph(recognition_text)

# Punishment section
add_body_paragraph("Penal Measures and Deterrence")
punishment_text = (
    "The punitive measures introduced are uncompromising, reflecting the state's zero-tolerance policy. Section 103(2) mandates that every individual "
    "convicted of this offense faces either imprisonment for life or the death penalty, along with a mandatory fine. This effectively places mob lynching "
    "on par with the most severe crimes against the state and humanity. Additionally, the legislation displays foresight by including Section 117(4), "
    "which covers instances of grievous harm that do not result in death. Offenders in such cases face up to seven years of incarceration. This dual-layered "
    "approach ensures that the machinery of justice is equipped to handle both fatal and non-fatal manifestations of mob violence."
)
add_body_paragraph(punishment_text)

# Conclusion
add_body_paragraph("Conclusion")
conclusion_text = (
    "In conclusion, the Bharatiya Nyaya Sanhita, 2023, stands as a bulwark against the normalization of mob justice. By replacing the antiquated IPC "
    "framework with modern, specific provisions, strict liability is established for hate crimes. The rigorous punishments prescribed serve as a stern "
    "warning that the rule of law will always supersede the rule of the mob. This legislative evolution is crucial for preserving the democratic ethos of "
    "India, ensuring that the diversity of the nation is protected by the iron hand of the law."
)
add_body_paragraph(conclusion_text)

# References
document.add_page_break()
ref_heading = document.add_paragraph()
ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = ref_heading.add_run("References")
set_font(run, bold=True, size=14)

def add_reference_with_link(text, url=None):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.0
    p_format.space_after = Pt(6)
    
    # Split text to isolate the part that should be linked if needed, 
    # but for simplicity, we can just append the link or make the whole citation a link if appropriate.
    # Here, we will write the text, and if a URL is provided, add a clickable "[Link]" at the end.
    
    run = p.add_run(text + " ")
    set_font(run, size=12)
    
    if url:
        add_hyperlink(p, "[View Source]", url)

add_reference_with_link("[1] Section 103, The Bharatiya Nyaya Sanhita, 2023 (Act No. 45 of 2023).", "https://www.indiacode.nic.in/bitstream/123456789/21033/1/the_bharatiya_nyaya_sanhita%2C_2023.pdf")
add_reference_with_link("[2] Ibid, Section 103(2).")
add_reference_with_link("[3] Section 117(4), The Bharatiya Nyaya Sanhita, 2023.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")
add_reference_with_link("[4] PRS Legislative Research, 'The Bharatiya Nyaya Sanhita, 2023 Bill Summary'.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")

# Save the document
filename = "A-15 Rakesh Mishra_Mob_Lynching_Assignment.docx"
document.save(filename)
print(f"Document saved as {filename}")
