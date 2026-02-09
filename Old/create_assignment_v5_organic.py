from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

document = Document()

# Set default style
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

def set_font(run, bold=False, italic=False, size=14):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# --- Cover Page ---
title_paragraph = document.add_paragraph()
title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_paragraph.add_run("\n\n\n\nASSIGNMENT\n")
set_font(run, bold=True, size=20)

run = title_paragraph.add_run("ON\n")
set_font(run, size=16)

run = title_paragraph.add_run("THE BHARATIYA NYAYA SANHITA 2023\n\n")
set_font(run, bold=True, size=18)

run = title_paragraph.add_run("Topic:\n")
set_font(run, size=16)

run = title_paragraph.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita\n\n\n")
set_font(run, bold=True, size=18)

# Student Details
details = document.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run("Submitted by:\n")
set_font(run, size=14)
run = details.add_run("Name: [Your Name Here]\n") 
set_font(run, bold=True, size=14)
run = details.add_run("Roll No: [Roll No]\n")
set_font(run, bold=True, size=14)
run = details.add_run("Div: [Division]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Submitted to:\n")
set_font(run, size=14)
run = details.add_run("[Professor Name]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Class:\n")
set_font(run, size=14)
run = details.add_run("Third Year BA LLB (Semester VI)\n(A.Y-2025-26)\n\n")
set_font(run, bold=True, size=14)

document.add_page_break()

# --- Main Content ---

heading = document.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita 2023")
set_font(run, bold=True, size=16)

def add_body_paragraph(text):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=14)
    return p

# ORGANIC / STUDENT-LIKE CONTENT

add_body_paragraph("Introduction")
intro_text = (
    "For decades, India struggled with a weird legal gap. We saw mobs attacking people, but there wasn't a specific law to punish the 'mob' itself. "
    "The police just used the normal murder sections from the old IPC (Indian Penal Code), which didn't really work well because pinning the exact "
    "blame on a chaotic crowd is a nightmare. But things changed effectively from July 2024. The new Bharatiya Nyaya Sanhita (BNS) came into the picture, "
    "and it finally calls out mob lynching for what it actually is. It is not just a riot; it is a specific kind of hate crime that needed its own rule."
)
add_body_paragraph(intro_text)

add_body_paragraph("Legal Recognition under Section 103(2)")
recognition_text = (
    "So, how does the new law handle it? Specifically, look at Section 103(2). It basically says if you have a group—five people or more—and they kill "
    "someone based on stuff like caste, language, race, or even personal beliefs, they are all guilty. It treats them as a single unit. This is the "
    "game-changer. Before this, people in a mob could claim they were just 'standing there' or didn't strike the fatal blow. The BNS removes that excuse. "
    "If you are part of that group acting together with the same intention, you are responsible for the murder. Period."
)
add_body_paragraph(recognition_text)

add_body_paragraph("Punishment and Deterrence")
punishment_text = (
    "The consequences are surprisingly serious now. The law doesn't go easy on this at all. Under Section 103(2), the punishment is either life imprisonment "
    "or the death penalty, plus a fine. It puts lynching right up there with the worst crimes possible. And it’s not just about killing, either. "
    "Section 117(4) covers cases where the victim survives but gets badly hurt (what the law calls 'grievous hurt'). Even then, the jail time can go "
    "up to seven years. So, whether the victim dies or gets maimed, the law now has a specific section to catch the attackers."
)
add_body_paragraph(punishment_text)

add_body_paragraph("Conclusion")
conclusion_text = (
    "Honestly, this was long overdue. Replacing the British-era IPC with the BNS shows that our legal system is finally waking up to modern problems. "
    "By explicitly threatening the death penalty and life jail terms, the message is super clear: mob rule is over. It gives the courts the right tools "
    "to actually punish these groups, instead of letting them slip away due to technicalities. It is a massive step for protecting people’s rights and "
    "stopping vigilantism."
)
add_body_paragraph(conclusion_text)

# References
document.add_page_break()
ref_heading = document.add_paragraph()
ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = ref_heading.add_run("References")
set_font(run, bold=True, size=14)

def add_reference_with_link(text, url=None):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.0
    p_format.space_after = Pt(6)
    
    run = p.add_run(text + " ")
    set_font(run, size=12)
    
    if url:
        add_hyperlink(p, "[View Source]", url)

add_reference_with_link("[1] Section 103, The Bharatiya Nyaya Sanhita, 2023 (Act No. 45 of 2023).", "https://www.indiacode.nic.in/bitstream/123456789/21033/1/the_bharatiya_nyaya_sanhita%2C_2023.pdf")
add_reference_with_link("[2] Ibid, Section 103(2).")
add_reference_with_link("[3] Section 117(4), The Bharatiya Nyaya Sanhita, 2023.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")
add_reference_with_link("[4] PRS Legislative Research, 'The Bharatiya Nyaya Sanhita, 2023 Bill Summary'.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")

# Save the document
filename = "A-15 Rakesh Mishra_Mob_Lynching_Assignment_v5_Organic.docx"
document.save(filename)
print(f"Document saved as {filename}")
