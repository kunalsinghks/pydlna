from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

document = Document()

# Set default style
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

def set_font(run, bold=False, italic=False, size=14):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# --- Cover Page ---
title_paragraph = document.add_paragraph()
title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_paragraph.add_run("\n\n\n\nASSIGNMENT\n")
set_font(run, bold=True, size=20)

run = title_paragraph.add_run("ON\n")
set_font(run, size=16)

run = title_paragraph.add_run("THE BHARATIYA NYAYA SANHITA 2023\n\n")
set_font(run, bold=True, size=18)

run = title_paragraph.add_run("Topic:\n")
set_font(run, size=16)

run = title_paragraph.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita\n\n\n")
set_font(run, bold=True, size=18)

# Student Details
details = document.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run("Submitted by:\n")
set_font(run, size=14)
run = details.add_run("Name: [Your Name Here]\n") 
set_font(run, bold=True, size=14)
run = details.add_run("Roll No: [Roll No]\n")
set_font(run, bold=True, size=14)
run = details.add_run("Div: [Division]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Submitted to:\n")
set_font(run, size=14)
run = details.add_run("[Professor Name]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Class:\n")
set_font(run, size=14)
run = details.add_run("Third Year BA LLB (Semester VI)\n(A.Y-2025-26)\n\n")
set_font(run, bold=True, size=14)

document.add_page_break()

# --- Main Content ---

heading = document.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita 2023")
set_font(run, bold=True, size=16)

def add_body_paragraph(text):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=14)
    return p

# HUMANIZED CONTENT

add_body_paragraph("Introduction")
intro_text = (
    "Mob lynching has been a big problem in India for years, but our old criminal laws, the Indian Penal Code (IPC), never really had a specific rule for it. "
    "Police usually had to use normal murder laws, which made it very hard to catch everyone involved in a big, angry crowd. This gap in the law let a lot "
    "of people get away with violence. Now, with the new Bharatiya Nyaya Sanhita (BNS) 2023 coming into effect in July 2024, this has finally changed. "
    "The new law actually recognizes that mob violence is different from a normal crime and needs to be handled with its own set of rules. It is a major "
    "step forward in acknowledging the reality of hate crimes in our society."
)
add_body_paragraph(intro_text)

add_body_paragraph("Legal Recognition under Section 103(2)")
recognition_text = (
    "The most important part of this new law is Section 103(2). It clearly says that if a group of five or more people work together to kill someone because "
    "of their race, caste, community, language, or even personal beliefs, it will be treated as mob lynching. This is a huge shift from the past. Under the "
    "old IPC, it was often difficult to prove who did what in a large crowd, and many attackers got away with it because of 'lack of evidence' against "
    "specific individuals. Now, the law uses the concept of 'acting in concert.' This means that every member of that unlawful assembly is equally "
    "responsible for the death. It basically stops people from hiding behind the crowd; if you are part of the mob that killed someone, you are considered "
    "guilty of the murder."
)
add_body_paragraph(recognition_text)

add_body_paragraph("Punishment and Deterrence")
punishment_text = (
    "The government knows that strong words aren't enough, so they have set very harsh punishments to scare potential offenders. If a person is found guilty "
    "under Section 103(2), they face the most severe penalties allowed in our system: imprisonment for life or even the death penalty, along with a fine. "
    "This places mob lynching on the same level as the worst crimes in the country. But the law goes further than just murder cases. It also looks at "
    "situations where the victim doesn't die but gets badly injured. Section 117(4) deals with 'grievous hurt' caused by a mob. In these situations, the "
    "guilty people can face jail time of up to seven years. This shows that the law is trying to cover all aspects of mob violence, ensuring nobody slips "
    "through the cracks."
)
add_body_paragraph(punishment_text)

add_body_paragraph("Conclusion")
conclusion_text = (
    "Overall, the Bharatiya Nyaya Sanhita is a much-needed update for our legal system. By clearly defining mob lynching and setting harsh punishments like "
    "the death penalty, it sends a clear message that mob rule won't be tolerated anymore. It gives the courts the power they need to punish these crimes "
    "properly without having to rely on vague interpretations of old laws. Hopefully, this strict stance will stop people from taking the law into their "
    "own hands and will help protect the rights and safety of every citizen, no matter who they are or where they come from."
)
add_body_paragraph(conclusion_text)

# References
document.add_page_break()
ref_heading = document.add_paragraph()
ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = ref_heading.add_run("References")
set_font(run, bold=True, size=14)

def add_reference_with_link(text, url=None):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.0
    p_format.space_after = Pt(6)
    
    run = p.add_run(text + " ")
    set_font(run, size=12)
    
    if url:
        add_hyperlink(p, "[View Source]", url)

add_reference_with_link("[1] Section 103, The Bharatiya Nyaya Sanhita, 2023 (Act No. 45 of 2023).", "https://width=device-width://www.indiacode.nic.in/bitstream/123456789/21033/1/the_bharatiya_nyaya_sanhita%2C_2023.pdf")
add_reference_with_link("[2] Ibid, Section 103(2).")
add_reference_with_link("[3] Section 117(4), The Bharatiya Nyaya Sanhita, 2023.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")
add_reference_with_link("[4] PRS Legislative Research, 'The Bharatiya Nyaya Sanhita, 2023 Bill Summary'.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")

# Save the document
filename = "A-15 Rakesh Mishra_Mob_Lynching_Assignment_v4_Human.docx"
document.save(filename)
print(f"Document saved as {filename}")
