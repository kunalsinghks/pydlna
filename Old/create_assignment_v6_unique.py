from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

document = Document()

# Set default style
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

def set_font(run, bold=False, italic=False, size=14):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# --- Cover Page ---
title_paragraph = document.add_paragraph()
title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_paragraph.add_run("\n\n\n\nASSIGNMENT\n")
set_font(run, bold=True, size=20)

run = title_paragraph.add_run("ON\n")
set_font(run, size=16)

run = title_paragraph.add_run("THE BHARATIYA NYAYA SANHITA 2023\n\n")
set_font(run, bold=True, size=18)

run = title_paragraph.add_run("Topic:\n")
set_font(run, size=16)

run = title_paragraph.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita\n\n\n")
set_font(run, bold=True, size=18)

# Student Details
details = document.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run("Submitted by:\n")
set_font(run, size=14)
run = details.add_run("Name: [Your Name Here]\n") 
set_font(run, bold=True, size=14)
run = details.add_run("Roll No: [Roll No]\n")
set_font(run, bold=True, size=14)
run = details.add_run("Div: [Division]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Submitted to:\n")
set_font(run, size=14)
run = details.add_run("[Professor Name]\n\n")
set_font(run, bold=True, size=14)

run = details.add_run("Class:\n")
set_font(run, size=14)
run = details.add_run("Third Year BA LLB (Semester VI)\n(A.Y-2025-26)\n\n")
set_font(run, bold=True, size=14)

document.add_page_break()

# --- Main Content ---

heading = document.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run("Mob Lynching: Legal Recognition and Punishment under Bharatiya Nyaya Sanhita 2023")
set_font(run, bold=True, size=16)

def add_body_paragraph(text):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=14)
    return p

# RADICALLY REWRITTEN CONTENT (LOW PLAGIARISM FOCUS)

add_body_paragraph("Introduction")
intro_text = (
    "In July 2024, India turned a new page in its judicial history by discarding the colonial-era Indian Penal Code (IPC) and enforcing the Bharatiya "
    "Nyaya Sanhita (BNS). This shift was not merely administrative; it represented a deeper societal need to confront crimes that previous laws ignored. "
    "For years, vigilante justice and mob violence operated in a legal vacuum, often treated no differently than ordinary street fights or riots. The BNS "
    "has finally closed this loop. By institutionalizing specific penalties for what we commonly call 'lynching,' the state has acknowledged that "
    "collective hatred is a distinct poison requiring a distinct antidote."
)
add_body_paragraph(intro_text)

add_body_paragraph("Legal Framework and Recognition")
recognition_text = (
    "The core of this reform lies in Section 103(2), a provision designed to pin accountability on faceless crowds. Previously, the defense strategies "
    "in mob violence cases relied heavily on the confusion of the moment—claiming it was impossible to know who struck the fatal blow. The new law disarms "
    "this argument. It explicitly states that when five or more individuals gather with a shared intent to kill based on prejudice—be it caste, community, "
    "or personal belief—they are all equally culpable. The law no longer sees a 'crowd'; it sees a coordinated unit of offenders, stripping away the "
    "defense of anonymity."
)
add_body_paragraph(recognition_text)

add_body_paragraph("Penal Measures and Deterrence")
punishment_text = (
    "To match the severity of the crime, the BNS has introduced uncompromising penalties. The statute mandates that anyone convicted under Section "
    "103(2) faces life imprisonment or the death penalty, in addition to a substantial fine. This effectively categorizes hate-driven mob murder as "
    "a capital offense, signaling zero tolerance from the judicial system. Furthermore, the legislation anticipates scenarios where victims survive. "
    "Section 117(4) specifically targets non-fatal mob attacks resulting in 'grievous hurt,' prescribing prison terms of up to seven years. This "
    "ensures a comprehensive legal net that catches all forms of group violence, regardless of the lethal outcome."
)
add_body_paragraph(punishment_text)

add_body_paragraph("Conclusion")
conclusion_text = (
    "Ultimately, the specific inclusion of mob lynching in the BNS is a victory for constitutional values over majoritarian impulses. It transforms "
    "the abstract promise of 'equality before law' into a concrete enforcement mechanism. By equating mob justice with the highest crimes known to "
    "the state, the new code sends a stern warning: the power to punish belongs solely to the judiciary, never to the street. This legal evolution "
    "is essential for safeguarding the dignity of every citizen against the tyranny of the crowd."
)
add_body_paragraph(conclusion_text)

# References
document.add_page_break()
ref_heading = document.add_paragraph()
ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = ref_heading.add_run("References")
set_font(run, bold=True, size=14)

def add_reference_with_link(text, url=None):
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing = 1.0
    p_format.space_after = Pt(6)
    
    run = p.add_run(text + " ")
    set_font(run, size=12)
    
    if url:
        add_hyperlink(p, "[View Source]", url)

add_reference_with_link("[1] Section 103, The Bharatiya Nyaya Sanhita, 2023 (Act No. 45 of 2023).", "https://www.indiacode.nic.in/bitstream/123456789/21033/1/the_bharatiya_nyaya_sanhita%2C_2023.pdf")
add_reference_with_link("[2] Ibid, Section 103(2).")
add_reference_with_link("[3] Section 117(4), The Bharatiya Nyaya Sanhita, 2023.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")
add_reference_with_link("[4] PRS Legislative Research, 'The Bharatiya Nyaya Sanhita, 2023 Bill Summary'.", "https://prsindia.org/billtrack/the-bharatiya-nyaya-sanhita-2023")

# Save the document
filename = "A-15 Rakesh Mishra_Mob_Lynching_Assignment_v6_Unique.docx"
document.save(filename)
print(f"Document saved as {filename}")
